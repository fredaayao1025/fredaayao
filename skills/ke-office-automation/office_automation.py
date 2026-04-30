#!/usr/bin/env python3
"""
办公自动化脚本 - 科特船长版
功能：Excel 合并/文件重命名/PDF 处理/Word 批量处理
"""

import os
import sys
import json
import shutil
from pathlib import Path
from datetime import datetime

def merge_excel_files(input_dir, output_file):
    """合并 Excel 文件"""
    print(f"📁 扫描目录：{input_dir}")
    
    input_path = Path(input_dir)
    if not input_path.exists():
        print(f"❌ 目录不存在：{input_dir}")
        return
    
    excel_files = list(input_path.glob("*.xlsx")) + list(input_path.glob("*.xls"))
    excel_files = sorted(excel_files)
    
    if not excel_files:
        print("❌ 未找到 Excel 文件")
        print("💡 提示：确保文件在指定目录中")
        return
    
    print(f"✅ 找到 {len(excel_files)} 个 Excel 文件:")
    for i, f in enumerate(excel_files, 1):
        print(f"   {i}. {f.name}")
    
    # 生成合并报告
    report = f"# Excel 合并报告\n\n"
    report += f"**扫描目录**: {input_dir}\n"
    report += f"**文件数量**: {len(excel_files)}\n"
    report += f"**扫描时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
    report += "## 文件列表\n\n"
    for f in excel_files:
        report += f"- {f.name}\n"
    
    # 尝试使用 pandas 合并
    try:
        import pandas as pd
        print("\n📊 使用 pandas 合并 Excel 文件...")
        
        all_data = []
        for f in excel_files:
            try:
                df = pd.read_excel(f)
                df['源文件'] = f.name
                all_data.append(df)
                print(f"   ✅ {f.name}: {len(df)} 行")
            except Exception as e:
                print(f"   ⚠️ {f.name}: 读取失败 - {e}")
        
        if all_data:
            merged_df = pd.concat(all_data, ignore_index=True)
            
            # 输出文件
            if output_file.endswith('.xlsx'):
                merged_df.to_excel(output_file, index=False)
            elif output_file.endswith('.csv'):
                merged_df.to_csv(output_file, index=False, encoding='utf-8-sig')
            else:
                output_file = output_file + '.xlsx'
                merged_df.to_excel(output_file, index=False)
            
            print(f"\n✅ 合并完成！")
            print(f"📄 输出文件：{output_file}")
            print(f"📊 总行数：{len(merged_df)}")
            
            report += f"\n## 合并结果\n\n"
            report += f"- 总行数：{len(merged_df)}\n"
            report += f"- 输出文件：{output_file}\n"
    
    except ImportError:
        print("\n⚠️  pandas 未安装，无法合并数据")
        print("💡 安装命令：pip install pandas openpyxl")
        report += f"\n## 提示\n\n需要安装 pandas 和 openpyxl 才能合并数据\n"
        report += "```\npip install pandas openpyxl\n```\n"
    
    # 保存报告
    report_file = Path(output_file).with_suffix('.md')
    report_file.write_text(report, encoding='utf-8')
    print(f"📄 报告已保存：{report_file}")

def rename_files(input_dir, pattern):
    """批量重命名文件"""
    print(f"📁 扫描目录：{input_dir}")
    
    input_path = Path(input_dir)
    if not input_path.exists():
        print(f"❌ 目录不存在：{input_dir}")
        return
    
    files = [f for f in input_path.iterdir() if f.is_file()]
    files = sorted(files, key=lambda f: f.stat().st_mtime)  # 按修改时间排序
    
    if not files:
        print("❌ 未找到文件")
        return
    
    print(f"✅ 找到 {len(files)} 个文件")
    
    date_str = datetime.now().strftime("%Y%m%d")
    renamed_count = 0
    
    print(f"\n📝 重命名预览:")
    for i, file in enumerate(files, 1):
        if "{original}" in pattern:
            new_name = pattern.replace("{original}", file.stem)
            new_name = new_name.replace("{date}", date_str)
        elif "{date}" in pattern:
            new_name = pattern.replace("{date}", date_str) + f"_{i:03d}"
        else:
            new_name = pattern + f"_{i:03d}"
        
        new_name = new_name + file.suffix
        new_path = file.parent / new_name
        
        if new_path != file:
            print(f"   {file.name} → {new_name}")
            renamed_count += 1
    
    # 确认执行
    print(f"\n✅ 将重命名 {renamed_count} 个文件")
    
    # 执行重命名
    for i, file in enumerate(files, 1):
        if "{original}" in pattern:
            new_name = pattern.replace("{original}", file.stem)
            new_name = new_name.replace("{date}", date_str)
        elif "{date}" in pattern:
            new_name = pattern.replace("{date}", date_str) + f"_{i:03d}"
        else:
            new_name = pattern + f"_{i:03d}"
        
        new_name = new_name + file.suffix
        new_path = file.parent / new_name
        
        if new_path != file:
            try:
                shutil.move(str(file), str(new_path))
            except Exception as e:
                print(f"   ❌ {file.name} 重命名失败：{e}")
    
    print(f"\n✅ 重命名完成！")

def merge_pdf_files(input_dir, output_file):
    """合并 PDF 文件"""
    print(f"📁 扫描目录：{input_dir}")
    
    input_path = Path(input_dir)
    if not input_path.exists():
        print(f"❌ 目录不存在：{input_dir}")
        return
    
    pdf_files = list(input_path.glob("*.pdf"))
    pdf_files = sorted(pdf_files)
    
    if not pdf_files:
        print("❌ 未找到 PDF 文件")
        return
    
    print(f"✅ 找到 {len(pdf_files)} 个 PDF 文件:")
    for i, f in enumerate(pdf_files, 1):
        print(f"   {i}. {f.name}")
    
    # 尝试使用 pypdf 合并
    try:
        from pypdf import PdfWriter, PdfReader
        print("\n📊 合并 PDF 文件...")
        
        merger = PdfWriter()
        
        for f in pdf_files:
            try:
                reader = PdfReader(f)
                merger.append(reader)
                print(f"   ✅ {f.name}: {len(reader.pages)} 页")
            except Exception as e:
                print(f"   ⚠️ {f.name}: 添加失败 - {e}")
        
        # 输出文件
        if not output_file.endswith('.pdf'):
            output_file = output_file + '.pdf'
        
        merger.write(output_file)
        merger.close()
        
        print(f"\n✅ 合并完成！")
        print(f"📄 输出文件：{output_file}")
    
    except ImportError:
        print("\n⚠️  pypdf 未安装，无法合并 PDF")
        print("💡 安装命令：pip install pypdf")
    
    # 生成报告
    report = f"# PDF 合并报告\n\n"
    report += f"**扫描目录**: {input_dir}\n"
    report += f"**文件数量**: {len(pdf_files)}\n\n"
    report += "## 文件列表\n\n"
    for f in pdf_files:
        report += f"- {f.name}\n"
    
    report_file = Path(output_file).with_suffix('.md')
    report_file.write_text(report, encoding='utf-8')
    print(f"📄 报告已保存：{report_file}")

def word_batch_replace(input_dir, old_text, new_text):
    """Word 批量替换文本"""
    print(f"📁 扫描目录：{input_dir}")
    
    input_path = Path(input_dir)
    if not input_path.exists():
        print(f"❌ 目录不存在：{input_dir}")
        return
    
    docx_files = list(input_path.glob("*.docx"))
    
    if not docx_files:
        print("❌ 未找到 Word 文件")
        return
    
    print(f"✅ 找到 {len(docx_files)} 个 Word 文件")
    print(f"📝 替换内容：'{old_text}' → '{new_text}'")
    
    # 尝试使用 python-docx
    try:
        from docx import Document
        print("\n📊 开始批量替换...")
        
        for f in docx_files:
            try:
                doc = Document(f)
                
                # 替换段落中的文本
                replaced = False
                for paragraph in doc.paragraphs:
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, new_text)
                        replaced = True
                
                # 替换表格中的文本
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if old_text in cell.text:
                                cell.text = cell.text.replace(old_text, new_text)
                                replaced = True
                
                if replaced:
                    doc.save(f)
                    print(f"   ✅ {f.name}: 已替换")
                else:
                    print(f"   ⚠️ {f.name}: 未找到匹配文本")
                
            except Exception as e:
                print(f"   ❌ {f.name}: 处理失败 - {e}")
        
        print(f"\n✅ 批量替换完成！")
    
    except ImportError:
        print("\n⚠️  python-docx 未安装，无法处理 Word")
        print("💡 安装命令：pip install python-docx")

def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("=" * 60)
        print("办公自动化脚本 - 科特船长版")
        print("=" * 60)
        print()
        print("用法：python office_automation.py <action> [options]")
        print()
        print("可用操作:")
        print("  excel-merge   合并 Excel 文件")
        print("  rename        批量重命名文件")
        print("  pdf-merge     合并 PDF 文件")
        print("  word-replace  Word 批量替换")
        print()
        print("示例:")
        print("  python office_automation.py excel-merge --input ./data --output ./merged.xlsx")
        print("  python office_automation.py rename --input ./files --pattern '{date}_{original}'")
        print("  python office_automation.py pdf-merge --input ./pdfs --output ./merged.pdf")
        print("  python office_automation.py word-replace --input ./docs --old '旧文本' --new '新文本'")
        print()
        return
    
    action = sys.argv[1]
    
    # 解析参数
    def get_arg(name, default=None):
        try:
            idx = sys.argv.index(f"--{name}")
            if idx + 1 < len(sys.argv):
                return sys.argv[idx + 1]
        except ValueError:
            pass
        return default
    
    if action == "excel-merge":
        input_dir = get_arg("input", "./data")
        output = get_arg("output", "./merged.xlsx")
        merge_excel_files(input_dir, output)
    
    elif action == "rename":
        input_dir = get_arg("input", "./files")
        pattern = get_arg("pattern", "{date}_{original}")
        rename_files(input_dir, pattern)
    
    elif action == "pdf-merge":
        input_dir = get_arg("input", "./pdfs")
        output = get_arg("output", "./merged.pdf")
        merge_pdf_files(input_dir, output)
    
    elif action == "word-replace":
        input_dir = get_arg("input", "./docs")
        old_text = get_arg("old", "")
        new_text = get_arg("new", "")
        if not old_text:
            print("❌ 错误：--old 参数必填")
            return
        word_batch_replace(input_dir, old_text, new_text)
    
    else:
        print(f"❌ 未知操作：{action}")
        print("运行脚本 without 参数查看帮助")

if __name__ == '__main__':
    main()
